"""
江西财经大学现代经济管理学院本科毕业论文 DOCX 自动排版工具（v4.2）

基于 v4，仅修复 add_heading_formatted：
  - 二三级标题段前段后 = 0（用 w:before="0" twips）
  - 一级标题段前段后 = 0.5行（用 w:beforeLines="50"）

依赖：pip install python-docx
用法：cd scripts && python thesis_formatter.py
"""

import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import config as cfg


# ============================================================
# 工具函数
# ============================================================

def set_run_font(run, cn_font, en_font, size, bold=False):
    run.font.size = size
    run.font.name = en_font
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), cn_font)


def clear_paragraph(paragraph):
    el = paragraph._element
    for child in list(el):
        if child.tag == qn("w:pPr"):
            continue
        el.remove(child)


def set_first_line_indent_chars(paragraph, chars=2.0):
    val = str(int(chars * 100))
    pPr = paragraph._element.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    if ind.get(qn("w:firstLine")) is not None:
        del ind.attrib[qn("w:firstLine")]
    ind.set(qn("w:firstLineChars"), val)


def set_line_spacing_exact(paragraph, pt_value):
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = pt_value


def set_line_spacing_multiple(paragraph, multiple):
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = multiple


# ── 段落 ──

def add_body_paragraph(doc, text, cn_font, en_font, size, bold=False,
                       alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                       indent_chars=None, line_spacing_pt=None,
                       space_before=Pt(0), space_after=Pt(0)):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = space_before
    pf.space_after = space_after
    if line_spacing_pt is not None:
        set_line_spacing_exact(p, line_spacing_pt)
    if indent_chars is not None:
        set_first_line_indent_chars(p, indent_chars)
    run = p.add_run(text)
    set_run_font(run, cn_font, en_font, size, bold)
    return p


def add_title_paragraph(doc, text, cn_font, en_font, size, bold=True,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        space_before=Pt(0), space_after=Pt(0)):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = space_before
    pf.space_after = space_after
    run = p.add_run(text)
    set_run_font(run, cn_font, en_font, size, bold)
    return p


def add_heading_formatted(doc, text, level, cn_font, en_font, size, bold=True,
                          alignment=WD_ALIGN_PARAGRAPH.LEFT,
                          line_spacing_multiple=1.5,
                          before_lines=0, after_lines=0):
    """
    格式化标题（绑定 Heading 样式供 TOC 识别）
    before_lines / after_lines: 1/100 行
      50 = 0.5 行（一级标题）
      0  = 0 行（二三级标题）

    ★ 核心逻辑：
      > 0 时用 w:beforeLines（行单位）
      = 0 时用 w:before="0"（twips 写死 0，彻底覆盖样式默认值）
    """
    h = doc.add_heading(level=level)
    h.alignment = alignment
    set_line_spacing_multiple(h, line_spacing_multiple)

    # 段前段后处理
    pPr = h._element.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)

    # 先清除所有旧属性
    for attr in ["w:before", "w:after", "w:beforeLines", "w:afterLines",
                 "w:beforeAutospacing", "w:afterAutospacing"]:
        if spacing.get(qn(attr)) is not None:
            del spacing.attrib[qn(attr)]

    # 段前
    if before_lines > 0:
        spacing.set(qn("w:beforeLines"), str(before_lines))
    else:
        spacing.set(qn("w:before"), "0")

    # 段后
    if after_lines > 0:
        spacing.set(qn("w:afterLines"), str(after_lines))
    else:
        spacing.set(qn("w:after"), "0")

    # 清除默认 run，写入自定义字体
    clear_paragraph(h)
    run = h.add_run(text)
    set_run_font(run, cn_font, en_font, size, bold)
    return h


# ── 页码 ──

def set_page_number_field(paragraph, cn_font=None, en_font=None, size=None):
    run1 = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run1._element.append(fld_begin)

    run2 = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run2._element.append(instr)

    run3 = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run3._element.append(fld_end)

    if en_font and size:
        for r in [run1, run2, run3]:
            set_run_font(r, cn_font or en_font, en_font, size)


def set_roman_page_number(section):
    sectPr = section._sectPr
    for old in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(old)
    pgNum = OxmlElement("w:pgNumType")
    pgNum.set(qn("w:fmt"), "upperRoman")
    pgNum.set(qn("w:start"), "1")
    sectPr.append(pgNum)

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    clear_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_page_number_field(p)


def set_arabic_page_number(section, restart=True):
    sectPr = section._sectPr
    for old in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(old)
    pgNum = OxmlElement("w:pgNumType")
    pgNum.set(qn("w:fmt"), "decimal")
    if restart:
        pgNum.set(qn("w:start"), "1")
    sectPr.append(pgNum)

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    clear_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_page_number_field(p, cfg.FONT_SONGTI, cfg.FONT_TIMES, cfg.SIZE_XIAOWU)


def set_arabic_page_number_continue(section):
    sectPr = section._sectPr
    for old in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(old)

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    clear_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_page_number_field(p, cfg.FONT_SONGTI, cfg.FONT_TIMES, cfg.SIZE_XIAOWU)


# ── 页眉 ──

def add_header_with_border(section, text, font_cn, font_en, font_size):
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    clear_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run(text)
    set_run_font(run, font_cn, font_en, font_size)

    pPr = p._element.get_or_add_pPr()
    old_bdr = pPr.find(qn("w:pBdr"))
    if old_bdr is not None:
        pPr.remove(old_bdr)
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def remove_header(section):
    header = section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        clear_paragraph(p)
        pPr = p._element.find(qn("w:pPr"))
        if pPr is not None:
            pBdr = pPr.find(qn("w:pBdr"))
            if pBdr is not None:
                pPr.remove(pBdr)


# ── 封面 ──

def remove_cell_borders(cell):
    tc_pr = cell._element.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for name in ["top", "bottom", "left", "right", "insideH", "insideV"]:
        old = borders.find(qn(f"w:{name}"))
        if old is not None:
            borders.remove(old)
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)


def set_cell_text(cell, text, cn_font, en_font, size, bold=False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, cn_font, en_font, size, bold)


# ── 目录 ──

def insert_toc_field(doc):
    p = doc.add_paragraph()

    run1 = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")
    run1._element.append(fld_begin)

    run2 = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._element.append(instr)

    run3 = p.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run3._element.append(fld_sep)

    run4 = p.add_run("（请在 Word 中按 Ctrl+A → F9 更新目录）")
    set_run_font(run4, cfg.FONT_SONGTI, cfg.FONT_TIMES, cfg.SIZE_WUHAO)

    run5 = p.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run5._element.append(fld_end)


def setup_heading_styles(doc):
    """配置 Heading 样式：目录字体 + 清除默认段前段后"""
    for level, font_size, bold in [
        (1, cfg.SIZE_SIHAO, True),
        (2, cfg.SIZE_XIAOSI, False),
        (3, cfg.SIZE_XIAOSI, False),
    ]:
        style_name = f"Heading {level}"
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = cfg.FONT_TIMES
            style.font.size = font_size
            style.font.bold = bold
            style.font.color.rgb = None
            rpr = style._element.get_or_add_rPr()
            rFonts = rpr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rpr.insert(0, rFonts)
            rFonts.set(qn("w:eastAsia"), cfg.FONT_SONGTI)

            # 清除样式中自带的段前段后
            pPr = style._element.find(qn("w:pPr"))
            if pPr is not None:
                spacing = pPr.find(qn("w:spacing"))
                if spacing is not None:
                    for attr in ["w:before", "w:after",
                                 "w:beforeAutospacing", "w:afterAutospacing",
                                 "w:beforeLines", "w:afterLines"]:
                        if spacing.get(qn(attr)) is not None:
                            del spacing.attrib[qn(attr)]
                    # 样式层面也写死 0，防止继承
                    spacing.set(qn("w:before"), "0")
                    spacing.set(qn("w:after"), "0")


# ── Section ──

def create_section(doc):
    sec = doc.add_section()
    sec.top_margin = cfg.MARGIN_TOP
    sec.bottom_margin = cfg.MARGIN_BOTTOM
    sec.left_margin = cfg.MARGIN_LEFT
    sec.right_margin = cfg.MARGIN_RIGHT
    sec.page_width = cfg.PAGE_WIDTH
    sec.page_height = cfg.PAGE_HEIGHT
    return sec


# ============================================================
# 各部分生成
# ============================================================

def build_cover(doc):
    for _ in range(4):
        doc.add_paragraph()

    add_title_paragraph(doc, cfg.SCHOOL_NAME,
                        cfg.FONT_HEITI, cfg.FONT_TIMES, Pt(26),
                        bold=True, space_after=Pt(2))
    add_title_paragraph(doc, cfg.COLLEGE_NAME,
                        cfg.FONT_HEITI, cfg.FONT_TIMES, Pt(18),
                        bold=True, space_after=Pt(6))
    add_title_paragraph(doc, cfg.THESIS_TYPE,
                        cfg.FONT_HEITI, cfg.FONT_TIMES, Pt(22),
                        bold=True, space_after=Pt(24))
    add_title_paragraph(doc, cfg.TITLE,
                        cfg.FONT_HEITI, cfg.FONT_TIMES, Pt(18),
                        bold=True, space_after=Pt(6))
    if cfg.SUBTITLE:
        add_title_paragraph(doc, cfg.SUBTITLE,
                            cfg.FONT_HEITI, cfg.FONT_TIMES, Pt(15),
                            bold=False, space_after=Pt(12))

    for _ in range(3):
        doc.add_paragraph()

    items = [
        ("学生姓名", cfg.STUDENT_NAME),
        ("学    号", cfg.STUDENT_ID),
        ("院    系", cfg.DEPARTMENT),
        ("专    业", cfg.MAJOR),
        ("指导教师", cfg.ADVISOR),
        ("完成日期", cfg.DATE),
    ]
    table = doc.add_table(rows=len(items), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(items):
        left_cell = table.cell(i, 0)
        right_cell = table.cell(i, 1)
        set_cell_text(left_cell, label + "：", cfg.FONT_SONGTI, cfg.FONT_TIMES,
                      Pt(14), bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        set_cell_text(right_cell, value, cfg.FONT_SONGTI, cfg.FONT_TIMES,
                      Pt(14), bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        for cell in [left_cell, right_cell]:
            remove_cell_borders(cell)
        row = table.rows[i]
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Pt(32)


def build_abstract_cn(doc):
    add_title_paragraph(doc, cfg.ABSTRACT_TITLE_CN,
                        cfg.ABSTRACT_TITLE_FONT_CN, cfg.FONT_TIMES,
                        cfg.ABSTRACT_TITLE_SIZE,
                        bold=True, space_after=Pt(12))
    add_body_paragraph(doc,
                       "在此填写中文摘要内容。摘要应简明扼要地概述论文的研究目的、方法、"
                       "主要结果和结论，一般300~500个汉字。",
                       cfg.ABSTRACT_BODY_FONT_CN, cfg.FONT_TIMES,
                       cfg.ABSTRACT_BODY_SIZE,
                       indent_chars=cfg.FIRST_LINE_INDENT_CHARS,
                       line_spacing_pt=cfg.ABSTRACT_LINE_SPACING)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    set_line_spacing_exact(p, cfg.ABSTRACT_LINE_SPACING)
    run_label = p.add_run(cfg.KEYWORD_LABEL_CN)
    set_run_font(run_label, cfg.KEYWORD_LABEL_FONT_CN, cfg.FONT_TIMES,
                 cfg.KEYWORD_LABEL_SIZE, bold=True)
    run_value = p.add_run("关键词1；关键词2；关键词3；关键词4；关键词5")
    set_run_font(run_value, cfg.KEYWORD_BODY_FONT_CN, cfg.FONT_TIMES,
                 cfg.ABSTRACT_BODY_SIZE)


def build_abstract_en(doc):
    add_title_paragraph(doc, cfg.ABSTRACT_TITLE_EN,
                        cfg.FONT_TIMES, cfg.FONT_TIMES,
                        cfg.ABSTRACT_TITLE_SIZE,
                        bold=True, space_after=Pt(12))
    add_body_paragraph(doc,
                       "Write your English abstract here. The abstract should briefly "
                       "summarize the purpose, methods, main results, and conclusions "
                       "of the thesis. Approximately 300 words.",
                       cfg.FONT_TIMES, cfg.FONT_TIMES,
                       cfg.ABSTRACT_BODY_SIZE,
                       indent_chars=cfg.FIRST_LINE_INDENT_CHARS,
                       line_spacing_pt=cfg.ABSTRACT_LINE_SPACING)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    set_line_spacing_exact(p, cfg.ABSTRACT_LINE_SPACING)
    run_label = p.add_run(cfg.KEYWORD_LABEL_EN)
    set_run_font(run_label, cfg.FONT_TIMES, cfg.FONT_TIMES,
                 cfg.KEYWORD_LABEL_SIZE, bold=True)
    run_value = p.add_run(" keyword1; keyword2; keyword3; keyword4; keyword5")
    set_run_font(run_value, cfg.FONT_TIMES, cfg.FONT_TIMES,
                 cfg.ABSTRACT_BODY_SIZE)


def build_toc(doc):
    add_title_paragraph(doc, cfg.TOC_TITLE,
                        cfg.TOC_TITLE_FONT, cfg.FONT_TIMES,
                        cfg.TOC_TITLE_SIZE,
                        bold=True, space_after=Pt(12))
    insert_toc_field(doc)


def build_body(doc):
    # 论文题目
    add_title_paragraph(doc, cfg.TITLE,
                        cfg.BODY_TITLE_FONT, cfg.FONT_TIMES,
                        cfg.BODY_TITLE_SIZE,
                        bold=cfg.BODY_TITLE_BOLD,
                        space_after=Pt(12))

    for ch in cfg.DEFAULT_CHAPTERS:
        # 一级标题：段前 0.5 行 (50)，段后 0.5 行 (50)
        add_heading_formatted(
            doc, f"{ch['number']} {ch['title']}", level=1,
            cn_font=cfg.HEADING1_FONT, en_font=cfg.FONT_TIMES,
            size=cfg.HEADING1_SIZE, bold=cfg.HEADING1_BOLD,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            line_spacing_multiple=cfg.HEADING1_LINE_SPACING,
            before_lines=cfg.HEADING1_SPACE_BEFORE_LINES,  # 50
            after_lines=cfg.HEADING1_SPACE_AFTER_LINES,    # 50
        )

        for sec in ch.get("sections", []):
            # 二级标题：段前 0，段后 0
            add_heading_formatted(
                doc, f"{sec['number']} {sec['title']}", level=2,
                cn_font=cfg.HEADING2_FONT, en_font=cfg.FONT_TIMES,
                size=cfg.HEADING2_SIZE, bold=cfg.HEADING2_BOLD,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                line_spacing_multiple=cfg.HEADING2_LINE_SPACING,
                before_lines=0,
                after_lines=0,
            )

            # 正文
            add_body_paragraph(
                doc, sec["body"],
                cfg.BODY_FONT_CN, cfg.BODY_FONT_EN, cfg.BODY_SIZE,
                indent_chars=cfg.FIRST_LINE_INDENT_CHARS,
                line_spacing_pt=cfg.BODY_LINE_SPACING,
            )


def build_references(doc):
    add_heading_formatted(
        doc, cfg.REF_TITLE, level=1,
        cn_font=cfg.FONT_HEITI, en_font=cfg.FONT_TIMES,
        size=cfg.ABSTRACT_TITLE_SIZE, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        before_lines=0, after_lines=0,
    )
    refs = [
        "[1] 作者. 文献标题[J]. 期刊名, 年份, 卷(期): 页码.",
        "[2] 作者. 书名[M]. 出版地: 出版社, 年份.",
        "[3] Author A, Author B. Title of Paper[J]. Journal Name, 2024, 1(1): 1-10.",
        "[4] Author. Book Title[M]. Publisher, 2023.",
    ]
    for ref in refs:
        add_body_paragraph(doc, ref,
                           cfg.REF_FONT_CN, cfg.REF_FONT_EN, cfg.REF_SIZE,
                           alignment=WD_ALIGN_PARAGRAPH.LEFT,
                           line_spacing_pt=cfg.REF_LINE_SPACING,
                           space_after=Pt(2))


def build_appendix(doc):
    add_heading_formatted(
        doc, "附录", level=1,
        cn_font=cfg.FONT_HEITI, en_font=cfg.FONT_TIMES,
        size=cfg.ABSTRACT_TITLE_SIZE, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        before_lines=0, after_lines=0,
    )
    add_body_paragraph(doc,
                       "在此放置附录内容，如核心代码、补充数据表格、公式推演等。"
                       "按"附录1"、"附录2"依次编号。如果没有附录，请删除此页及目录中的对应条目。",
                       cfg.BODY_FONT_CN, cfg.BODY_FONT_EN, cfg.BODY_SIZE,
                       indent_chars=cfg.FIRST_LINE_INDENT_CHARS,
                       line_spacing_pt=cfg.BODY_LINE_SPACING)


def build_thanks(doc):
    add_heading_formatted(
        doc, cfg.THANKS_TITLE, level=1,
        cn_font=cfg.FONT_HEITI, en_font=cfg.FONT_TIMES,
        size=cfg.ABSTRACT_TITLE_SIZE, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        before_lines=0, after_lines=0,
    )
    add_body_paragraph(doc,
                       "在此填写致谢内容。感谢指导老师、同学、家人等对论文完成过程中的帮助与支持。",
                       cfg.THANKS_FONT, cfg.FONT_TIMES, cfg.THANKS_SIZE,
                       indent_chars=cfg.FIRST_LINE_INDENT_CHARS,
                       line_spacing_pt=cfg.THANKS_LINE_SPACING)


# ============================================================
# 主流程
# ============================================================

def create_thesis():
    doc = Document()

    # 全局页面
    sec0 = doc.sections[0]
    sec0.top_margin = cfg.MARGIN_TOP
    sec0.bottom_margin = cfg.MARGIN_BOTTOM
    sec0.left_margin = cfg.MARGIN_LEFT
    sec0.right_margin = cfg.MARGIN_RIGHT
    sec0.page_width = cfg.PAGE_WIDTH
    sec0.page_height = cfg.PAGE_HEIGHT

    # Normal 样式
    style = doc.styles["Normal"]
    style.font.name = cfg.FONT_TIMES
    style.font.size = cfg.BODY_SIZE
    rpr = style._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), cfg.FONT_SONGTI)

    # Heading 样式
    setup_heading_styles(doc)

    # ===== 1. 封面 =====
    build_cover(doc)
    doc.add_page_break()

    # ===== 2. 中文摘要（罗马数字页码，无页眉）=====
    sec_abs = create_section(doc)
    remove_header(sec_abs)
    set_roman_page_number(sec_abs)

    build_abstract_cn(doc)
    doc.add_page_break()

    # ===== 3. 英文摘要 =====
    build_abstract_en(doc)
    doc.add_page_break()

    # ===== 4. 目录 =====
    build_toc(doc)
    doc.add_page_break()

    # ===== 5. 正文（页眉 + 阿拉伯页码）=====
    sec_body = create_section(doc)
    add_header_with_border(sec_body, cfg.HEADER_TEXT,
                           cfg.HEADER_FONT, cfg.FONT_TIMES,
                           cfg.HEADER_FONT_SIZE)
    set_arabic_page_number(sec_body)

    build_body(doc)
    doc.add_page_break()

    # ===== 6. 参考文献（无页眉，页码续编）=====
    sec_ref = create_section(doc)
    remove_header(sec_ref)
    set_arabic_page_number_continue(sec_ref)

    build_references(doc)
    doc.add_page_break()

    # ===== 7. 附录（无页眉，页码续编）=====
    sec_app = create_section(doc)
    remove_header(sec_app)
    set_arabic_page_number_continue(sec_app)

    build_appendix(doc)
    doc.add_page_break()

    # ===== 8. 致谢（无页眉，页码续编）=====
    sec_thx = create_section(doc)
    remove_header(sec_thx)
    set_arabic_page_number_continue(sec_thx)

    build_thanks(doc)

    # 保存
    output_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(output_dir, "thesis_output.docx")
    doc.save(output_path)
    print(f"✅ 毕业论文已生成：{output_path}")
    print()
    print("📌 后续操作：")
    print("   1. 用 Microsoft Word 打开 thesis_output.docx")
    print("   2. Word 可能提示"是否更新域" → 点击"是"")
    print("   3. 或者手动：Ctrl+A 全选 → F9 更新所有域")
    print("   4. 检查目录、页眉、页码是否正确")


if __name__ == "__main__":
    create_thesis()
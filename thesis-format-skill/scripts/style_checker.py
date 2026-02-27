"""
江西财经大学本科毕业论文格式检查工具
用法：python style_checker.py <your_thesis.docx>
"""

import sys
import os

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

import config as cfg


class StyleChecker:
    def __init__(self, filepath):
        self.filepath = filepath
        self.doc = Document(filepath)
        self.issues = []

    def _add(self, location, expected, actual, field):
        self.issues.append({
            "location": location, "field": field,
            "expected": expected, "actual": actual,
        })

    def check_margins(self):
        for i, sec in enumerate(self.doc.sections):
            loc = f"Section {i + 1}"
            checks = [
                (sec.top_margin, cfg.MARGIN_TOP, "上边距"),
                (sec.bottom_margin, cfg.MARGIN_BOTTOM, "下边距"),
                (sec.left_margin, cfg.MARGIN_LEFT, "左边距"),
                (sec.right_margin, cfg.MARGIN_RIGHT, "右边距"),
            ]
            for actual, expected, name in checks:
                if actual != expected:
                    self._add(loc, str(expected), str(actual), name)

    def check_fonts(self):
        for i, para in enumerate(self.doc.paragraphs):
            if para.style.name.startswith("Heading") or not para.text.strip():
                continue
            for run in para.runs:
                preview = run.text[:20] if run.text else ""
                loc = f"段落 {i + 1} ('{preview}…')"

                if run.font.size and run.font.size != cfg.BODY_SIZE:
                    self._add(loc, str(cfg.BODY_SIZE), str(run.font.size), "字号")
                if run.font.name and run.font.name != cfg.FONT_TIMES:
                    self._add(loc, cfg.FONT_TIMES, run.font.name, "英文字体")

                rpr = run._element.find(qn("w:rPr"))
                if rpr is not None:
                    rf = rpr.find(qn("w:rFonts"))
                    if rf is not None:
                        ea = rf.get(qn("w:eastAsia"))
                        ok = [cfg.FONT_SONGTI, cfg.FONT_HEITI, cfg.FONT_KAITI]
                        if ea and ea not in ok:
                            self._add(loc, "/".join(ok), ea, "中文字体")

    def check_indent(self):
        for i, para in enumerate(self.doc.paragraphs):
            if para.style.name.startswith("Heading") or not para.text.strip():
                continue
            pPr = para._element.find(qn("w:pPr"))
            if pPr is not None:
                ind = pPr.find(qn("w:ind"))
                if ind is not None:
                    fl = ind.get(qn("w:firstLine"))
                    fc = ind.get(qn("w:firstLineChars"))
                    if fl and not fc:
                        preview = para.text[:20] if para.text else ""
                        self._add(f"段落 {i + 1} ('{preview}…')",
                                  "w:firstLineChars=200",
                                  f"w:firstLine={fl} (固定长度)",
                                  "首行缩进方式")

    def check_heading1_spacing(self):
        for i, para in enumerate(self.doc.paragraphs):
            if para.style.name != "Heading 1":
                continue
            pPr = para._element.find(qn("w:pPr"))
            if pPr is not None:
                spacing = pPr.find(qn("w:spacing"))
                if spacing is not None:
                    bl = spacing.get(qn("w:beforeLines"))
                    al = spacing.get(qn("w:afterLines"))
                    loc = f"一级标题 ('{para.text[:20]}…')"
                    if bl != "50":
                        self._add(loc, "beforeLines=50 (0.5行)",
                                  f"beforeLines={bl}", "段前")
                    if al != "50":
                        self._add(loc, "afterLines=50 (0.5行)",
                                  f"afterLines={al}", "段后")

    def check_headers(self):
        for i, sec in enumerate(self.doc.sections):
            if i == 0:
                continue
            header = sec.header
            if header.paragraphs:
                text = header.paragraphs[0].text.strip()
                if text and text != cfg.HEADER_TEXT:
                    self._add(f"Section {i + 1} 页眉",
                              cfg.HEADER_TEXT, text, "页眉文字")

    def run_all(self):
        print(f"🔍 正在检查：{self.filepath}")
        print("=" * 60)

        self.check_margins()
        self.check_fonts()
        self.check_indent()
        self.check_heading1_spacing()
        self.check_headers()

        if not self.issues:
            print("✅ 未发现格式问题！")
        else:
            print(f"⚠️  发现 {len(self.issues)} 个问题：\n")
            for i, iss in enumerate(self.issues, 1):
                print(f"  {i}. [{iss['field']}] {iss['location']}")
                print(f"     期望: {iss['expected']}")
                print(f"     实际: {iss['actual']}\n")

        print("=" * 60)
        print(f"共 {len(self.issues)} 个问题。")
        return self.issues


def main():
    if len(sys.argv) < 2:
        print("用法: python style_checker.py <your_thesis.docx>")
        sys.exit(1)
    checker = StyleChecker(sys